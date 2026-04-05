import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_cell_bg(cell, hex_color: str):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def create_docx(md_path, docx_path):
    doc = Document()
    
    for section in doc.sections:
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    current_table = None
    table_headers = []
    
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
            
        if line.startswith("---"):
            continue

        if line.startswith("# "):
            h = doc.add_heading(line[2:].strip(), level=1)
            for run in h.runs:
                run.font.color.rgb = RGBColor.from_string("1F5C99")
        elif line.startswith("## "):
            h = doc.add_heading(line[3:].strip(), level=2)
            for run in h.runs:
                run.font.color.rgb = RGBColor.from_string("1F5C99")
        elif line.startswith("### "):
            h = doc.add_heading(line[4:].strip(), level=3)
            for run in h.runs:
                run.font.color.rgb = RGBColor.from_string("2E6DA4")
        elif line.startswith("> "):
            p = doc.add_paragraph()
            run = p.add_run(line[2:].strip())
            run.font.italic = True
            p.paragraph_format.left_indent = Inches(0.5)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:].strip(), style='List Bullet')
        elif line.startswith("|"):
            if "---" in line:
                continue
            cols = [c.strip() for c in line.strip("|").split("|")]
            if not current_table:
                current_table = doc.add_table(rows=1, cols=len(cols))
                current_table.style = 'Table Grid'
                hdr_cells = current_table.rows[0].cells
                for i, text in enumerate(cols):
                    hdr_cells[i].text = text
                    hdr_cells[i].paragraphs[0].runs[0].bold = True
                    set_cell_bg(hdr_cells[i], "1F5C99")
                    hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            else:
                row_cells = current_table.add_row().cells
                for i, text in enumerate(cols):
                    if i < len(row_cells):
                        row_cells[i].text = text
        else:
            current_table = None
            doc.add_paragraph(line)

    doc.save(docx_path)
    print(f"Saved {docx_path}")

md_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Client_Tutorial_Guide.md")
docx_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Client_Tutorial_Guide.docx")

if os.path.exists(md_file):
    create_docx(md_file, docx_file)
else:
    print(f"Error: {md_file} not found!")

