"""
LittleSteps International Pre School — Documentation Generator
Generates a professional Documentation.docx for the Django project.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import os

# ──────────────────────────────────────────────
# Helper utilities
# ──────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Fill a table cell with a solid background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_page_break(doc):
    doc.add_page_break()

def add_heading(doc, text, level=1, color_hex="1F5C99"):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    return h

def add_paragraph(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)

def add_code_block(doc, code: str):
    """Render a plain-text code block using a shaded paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    # shaded background via XML
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'EEF2F7')
    pPr.append(shd)


# ──────────────────────────────────────────────
# Document build
# ──────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ══════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════

doc.add_paragraph()   # top spacer
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_para.add_run("LITTLE STEPS INTERNATIONAL PRE SCHOOL")
title_run.bold = True
title_run.font.size = Pt(26)
title_run.font.color.rgb = RGBColor(0x1F, 0x5C, 0x99)

subtitle_para = doc.add_paragraph()
subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle_para.add_run("Web Portal — Project Documentation")
subtitle_run.font.size = Pt(16)
subtitle_run.font.color.rgb = RGBColor(0x55, 0x77, 0xAA)

doc.add_paragraph()
doc.add_paragraph()

meta_lines = [
    ("Project Name:", "Little Steps International Pre School Web Portal"),
    ("Framework:",    "Django 6.0.3  (Python)"),
    ("Database:",     "PostgreSQL"),
    ("Version:",      "1.0.0"),
    ("Date:",         datetime.date.today().strftime("%B %d, %Y")),
    ("Author:",       "Development Team"),
]
for label, value in meta_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f"{label}  ")
    r1.bold = True
    r1.font.size = Pt(12)
    r2 = p.add_run(value)
    r2.font.size = Pt(12)

add_page_break(doc)

# ══════════════════════════════════════════════
# TABLE OF CONTENTS  (manual)
# ══════════════════════════════════════════════

add_heading(doc, "Table of Contents", level=1)

toc_entries = [
    ("1", "Project Overview"),
    ("2", "Technology Stack"),
    ("3", "Project Structure"),
    ("4", "Data Models"),
    ("5", "URL Routes"),
    ("6", "Views & Business Logic"),
    ("7", "Admin Panel"),
    ("8", "Templates"),
    ("9", "Static & Media Files"),
    ("10","Email Notification System"),
    ("11","Installation & Setup Guide"),
    ("12","Running the Application"),
    ("13","Features Summary"),
    ("14","Future Enhancements"),
]
for num, title in toc_entries:
    p = doc.add_paragraph()
    r = p.add_run(f"   {num}.   {title}")
    r.font.size = Pt(12)

add_page_break(doc)

# ══════════════════════════════════════════════
# 1. PROJECT OVERVIEW
# ══════════════════════════════════════════════

add_heading(doc, "1. Project Overview", level=1)
add_paragraph(doc,
    "Little Steps International Pre School Web Portal is a full-stack web application built with "
    "the Django framework. It serves as the digital face of the school and provides an integrated "
    "management system for school administrators, parents, and prospective families.",
    size=11
)
doc.add_paragraph()
add_paragraph(doc, "Core objectives of the portal:", bold=True)
bullets = [
    "Publish school notices and activity updates in real time.",
    "Showcase a categorised photo and video gallery.",
    "Accept and track online admission applications.",
    "Display an event calendar with upcoming and past events.",
    "Facilitate parent/guardian enquiries via the contact form.",
    "Provide an integrated admin back-end for school staff.",
    "Send automated email notifications for admissions and enquiries.",
]
for b in bullets:
    add_bullet(doc, b)

add_page_break(doc)

# ══════════════════════════════════════════════
# 2. TECHNOLOGY STACK
# ══════════════════════════════════════════════

add_heading(doc, "2. Technology Stack", level=1)

tech_table = doc.add_table(rows=1, cols=3)
tech_table.style = 'Table Grid'
tech_table.alignment = WD_TABLE_ALIGNMENT.LEFT

hdr_cells = tech_table.rows[0].cells
for i, text in enumerate(["Category", "Technology", "Version / Notes"]):
    hdr_cells[i].text = text
    hdr_cells[i].paragraphs[0].runs[0].bold = True
    set_cell_bg(hdr_cells[i], "1F5C99")
    hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

rows = [
    ("Backend Framework", "Django",           "6.0.3"),
    ("Language",          "Python",           "3.x"),
    ("Database",          "PostgreSQL",       "Production-grade relational DB"),
    ("Frontend",          "HTML5 / CSS3 / JS","Vanilla, custom CSS"),
    ("Media Storage",     "Cloudinary / S3",  "MEDIA_ROOT for user uploads"),
    ("Static Files",      "Whitenoise",       "CompressedManifestStaticFilesStorage"),
    ("Email Backend",     "SMTP (Production)","django.core.mail.backends.smtp"),
    ("Version Control",   "Git",              ".git repository"),
    ("Server",            "Gunicorn",         "Production WSGI HTTP Server"),
]
for row_data in rows:
    row = tech_table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

doc.add_paragraph()
add_page_break(doc)

# ══════════════════════════════════════════════
# 3. PROJECT STRUCTURE
# ══════════════════════════════════════════════

add_heading(doc, "3. Project Structure", level=1)
add_paragraph(doc, "Directory layout of the project root:", size=11)
doc.add_paragraph()

add_code_block(doc,
"""\
LittleSteps/
├── .env                       # Environment variables (IGNORED IN GIT)
├── manage.py                  # Django management entry-point
├── little_steps_web/          # Django project configuration package
│   ├── settings.py            # Production settings & security headers
│   ├── urls.py
│   ├── asgi.py
│   └── wsgi.py
├── school/                    # Main Django application
│   ├── models.py              # Database models
│   ├── views.py               # View functions
│   ├── urls.py                # URL patterns
│   ├── admin.py               # Admin configuration
│   ├── forms.py               # Django forms
│   └── migrations/            # Database migrations
├── templates/                 # HTML templates
│   ├── base.html
│   ├── home.html
│   ├── about.html
│   ├── gallery.html
│   ├── events.html
│   ├── updates.html
│   ├── admissions.html
│   └── contact.html
├── static/                    # Static assets
│   ├── css/                   # Stylesheet files
│   └── img/                   # Image assets
├── media/                     # User-uploaded files (runtime)
│   ├── gallery/
│   ├── events/
│   ├── activities/
│   ├── students/
│   └── teachers/
└── venv/                      # Python virtual environment\
""")

add_page_break(doc)

# ══════════════════════════════════════════════
# 4. DATA MODELS
# ══════════════════════════════════════════════

add_heading(doc, "4. Data Models", level=1)
add_paragraph(doc,
    "All data models are defined in school/models.py. They are managed "
    "through Django's ORM and accessed via the admin panel or views.",
    size=11
)
doc.add_paragraph()

models_info = [
    (
        "Notice",
        "Stores school announcements and notices.",
        [("title",       "CharField(200)",   "Heading of the notice"),
         ("content",     "TextField",        "Body text of the notice"),
         ("date_posted", "DateTimeField",    "Auto-set on creation"),
         ("is_active",   "BooleanField",     "Controls visibility on site")],
    ),
    (
        "GalleryCategory",
        "Categorises gallery images (e.g. Sports Day, Annual Day).",
        [("name", "CharField(100)", "Category label")],
    ),
    (
        "GalleryImage",
        "Individual images linked to a GalleryCategory.",
        [("category",   "ForeignKey(GalleryCategory)", "Parent category"),
         ("image",      "ImageField",                  "Image file (upload_to='gallery/')"),
         ("caption",    "CharField(200)",              "Optional caption"),
         ("created_at", "DateTimeField",               "Auto-set on creation")],
    ),
    (
        "AdmissionApplication",
        "Records online admission applications submitted by parents.",
        [("child_name",    "CharField(150)", "Full name of child"),
         ("date_of_birth", "DateField",      "Child's date of birth"),
         ("parent_name",   "CharField(150)", "Parent / guardian name"),
         ("email",         "EmailField",     "Contact email"),
         ("phone",         "CharField(15)",  "Contact phone number"),
         ("program",       "CharField",      "Chosen program (choices-based)"),
         ("message",       "TextField",      "Additional message"),
         ("submitted_at",  "DateTimeField",  "Auto-set on submission")],
    ),
    (
        "Event",
        "School events displayed on the Events page.",
        [("title",       "CharField(200)", "Event name"),
         ("description", "TextField",      "Event details"),
         ("date",        "DateField",      "Event date"),
         ("time",        "TimeField",      "Event time (optional)"),
         ("image",       "ImageField",     "Event banner image (optional)"),
         ("is_past",     "Property",       "True if event date < today")],
    ),
    (
        "Enquiry",
        "Contact/enquiry submissions from the contact page.",
        [("name",         "CharField(150)", "Sender name"),
         ("email",        "EmailField",     "Sender email"),
         ("phone",        "CharField(15)",  "Sender phone"),
         ("message",      "TextField",      "Enquiry message"),
         ("submitted_at", "DateTimeField",  "Auto-set on creation")],
    ),
    (
        "ActivityUpdate",
        "Daily activity blog posts managed by the admin.",
        [("title",   "CharField(200)", "Update heading"),
         ("content", "TextField",      "Update body"),
         ("image",   "ImageField",     "Activity photo"),
         ("date",    "DateField",      "Activity date (default today)")],
    ),
    (
        "VideoEmbed",
        "YouTube video embeds displayed on the home page.",
        [("title",       "CharField(200)", "Video title"),
         ("youtube_url", "URLField",       "Full YouTube URL"),
         ("is_active",   "BooleanField",   "Toggle visibility"),
         ("get_embed_url","Method",        "Extracts and returns embeddable URL")],
    ),
    (
        "Student",
        "Enrolled student records managed by the admin.",
        [("name",             "CharField(150)", "Student full name"),
         ("age",              "IntegerField",   "Student age"),
         ("program_enrolled", "CharField(200)", "Programme name"),
         ("parent_name",      "CharField(150)", "Parent full name"),
         ("parent_email",     "EmailField",     "Parent email (used for notifications)"),
         ("parent_phone",     "CharField(20)",  "Parent phone"),
         ("enrolled_date",    "DateField",      "Date of enrolment"),
         ("photo",            "ImageField",     "Student photo (optional)")],
    ),
    (
        "Teacher",
        "Teaching staff profiles shown on the About page.",
        [("name",        "CharField(150)", "Teacher full name"),
         ("designation", "CharField(150)", "Job title"),
         ("bio",         "TextField",      "Biography (optional)"),
         ("image",       "ImageField",     "Profile photo")],
    ),
]

for model_name, description, fields in models_info:
    add_heading(doc, f"4.{models_info.index((model_name, description, fields)) + 1}  {model_name}", level=2, color_hex="2E6DA4")
    add_paragraph(doc, description, size=11)
    doc.add_paragraph()

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    for i, h in enumerate(["Field / Method", "Type", "Description"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        set_cell_bg(hdr[i], "2E6DA4")
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for field, ftype, desc in fields:
        r = tbl.add_row().cells
        r[0].text = field
        r[1].text = ftype
        r[2].text = desc
    doc.add_paragraph()

add_page_break(doc)

# ══════════════════════════════════════════════
# 5. URL ROUTES
# ══════════════════════════════════════════════

add_heading(doc, "5. URL Routes", level=1)
add_paragraph(doc,
    "URL configuration is split between the project-level urls.py and the school app urls.py.",
    size=11
)
doc.add_paragraph()

url_tbl = doc.add_table(rows=1, cols=4)
url_tbl.style = 'Table Grid'
hdr = url_tbl.rows[0].cells
for i, h in enumerate(["URL Pattern", "View Function", "Name", "Description"]):
    hdr[i].text = h
    hdr[i].paragraphs[0].runs[0].bold = True
    set_cell_bg(hdr[i], "1F5C99")
    hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

url_rows = [
    ("/",            "home",       "home",       "Landing / home page"),
    ("/about/",      "about",      "about",      "About school & teachers"),
    ("/gallery/",    "gallery",    "gallery",    "Photo gallery"),
    ("/events/",     "events",     "events",     "Upcoming & past events"),
    ("/updates/",    "updates",    "updates",    "Daily activity updates"),
    ("/admissions/", "admissions", "admissions", "Online admission form"),
    ("/contact/",    "contact",    "contact",    "Contact/enquiry form"),
    ("/admin/",      "Django admin","—",         "Admin back-end"),
    ("/media/<path>","Static serve","—",         "Served media files (dev only)"),
]
for row_data in url_rows:
    row = url_tbl.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

add_page_break(doc)

# ══════════════════════════════════════════════
# 6. VIEWS & BUSINESS LOGIC
# ══════════════════════════════════════════════

add_heading(doc, "6. Views & Business Logic", level=1)

views_data = [
    ("home(request)",
     "Fetches up to 5 active notices, 3 latest activity updates, 3 upcoming events, "
     "and 2 active videos. Passes them to home.html."),
    ("about(request)",
     "Retrieves all Teacher objects and passes them to about.html for the staff profile section."),
    ("gallery(request)",
     "Loads all GalleryCategory and GalleryImage objects, enabling filtered display."),
    ("events(request)",
     "Splits events into upcoming (date >= today) and past (date < today, limited to 5). "
     "Displays them in distinct sections."),
    ("updates(request)",
     "Loads all ActivityUpdate records ordered by most recent date first."),
    ("admissions(request)",
     "GET: renders the admission form.\n"
     "POST: validates submitted data, creates an AdmissionApplication record, "
     "sends an email notification to the school, and shows a success message."),
    ("contact(request)",
     "GET: renders the contact form.\n"
     "POST: creates an Enquiry record, sends an email notification, "
     "and shows a success message."),
]

for func_name, desc in views_data:
    add_heading(doc, func_name, level=2, color_hex="2E6DA4")
    add_paragraph(doc, desc, size=11)
    doc.add_paragraph()

add_page_break(doc)

# ══════════════════════════════════════════════
# 7. ADMIN PANEL
# ══════════════════════════════════════════════

add_heading(doc, "7. Admin Panel", level=1)
add_paragraph(doc,
    "All models are registered in school/admin.py with custom configurations "
    "to provide a rich management interface.",
    size=11
)
doc.add_paragraph()

admin_tbl = doc.add_table(rows=1, cols=3)
admin_tbl.style = 'Table Grid'
hdr = admin_tbl.rows[0].cells
for i, h in enumerate(["Model Admin Class", "List Display", "Features"]):
    hdr[i].text = h
    hdr[i].paragraphs[0].runs[0].bold = True
    set_cell_bg(hdr[i], "1F5C99")
    hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

admin_rows = [
    ("NoticeAdmin",             "title, date_posted, is_active",          "Filter, search, bulk email action"),
    ("GalleryCategoryAdmin",    "name",                                    "Inline image upload (TabularInline)"),
    ("GalleryImageAdmin",       "caption, category, created_at",           "Filter by category"),
    ("AdmissionApplicationAdmin","child_name, program, parent_name, phone","Filter by program & date"),
    ("EventAdmin",              "title, date, time",                       "Filter by date, search"),
    ("EnquiryAdmin",            "name, email, phone, submitted_at",        "Filter by submitted date"),
    ("ActivityUpdateAdmin",     "title, date",                             "Filter by date, search"),
    ("VideoEmbedAdmin",         "title, is_active",                        "Toggle active status"),
    ("StudentAdmin",            "name, age, program_enrolled, parent_name","Filter by program & enrolled date"),
    ("TeacherAdmin",            "name, designation",                       "Search by name & designation"),
]
for row_data in admin_rows:
    row = admin_tbl.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

doc.add_paragraph()
add_heading(doc, "7.1  Bulk Email Action", level=2, color_hex="2E6DA4")
add_paragraph(doc,
    "A custom admin action send_notification_email is available on Notice and Event models. "
    "When triggered it sends individual notification emails to all parent emails stored in "
    "the Student model. The action reports how many notifications were sent successfully.",
    size=11
)

add_page_break(doc)

# ══════════════════════════════════════════════
# 8. TEMPLATES
# ══════════════════════════════════════════════

add_heading(doc, "8. Templates", level=1)
add_paragraph(doc,
    "All templates reside in the top-level templates/ directory. "
    "They follow a base → child inheritance pattern via base.html.",
    size=11
)
doc.add_paragraph()

tmpl_tbl = doc.add_table(rows=1, cols=3)
tmpl_tbl.style = 'Table Grid'
hdr = tmpl_tbl.rows[0].cells
for i, h in enumerate(["Template File", "URL", "Key Content"]):
    hdr[i].text = h
    hdr[i].paragraphs[0].runs[0].bold = True
    set_cell_bg(hdr[i], "1F5C99")
    hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

tmpl_rows = [
    ("base.html",       "—",             "Navigation bar, footer, CSS/JS link imports, message alerts"),
    ("home.html",       "/",             "Hero section, notices marquee, activity updates, events, videos, testimonials"),
    ("about.html",      "/about/",       "School mission/vision, programs, teacher profiles"),
    ("gallery.html",    "/gallery/",     "Category-filtered photo gallery"),
    ("events.html",     "/events/",      "Upcoming and past events listing"),
    ("updates.html",    "/updates/",     "Daily activity blog cards"),
    ("admissions.html", "/admissions/",  "Online admission application form"),
    ("contact.html",    "/contact/",     "Contact / enquiry form with school info"),
]
for row_data in tmpl_rows:
    row = tmpl_tbl.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

add_page_break(doc)

# ══════════════════════════════════════════════
# 9. STATIC & MEDIA FILES
# ══════════════════════════════════════════════

add_heading(doc, "9. Static & Media Files", level=1)

add_heading(doc, "9.1  Static Files", level=2, color_hex="2E6DA4")
add_paragraph(doc,
    "Static files (CSS, JavaScript, images) are stored in the static/ directory. "
    "The setting STATICFILES_DIRS points Django to static/ during development. "
    "Run python manage.py collectstatic before deploying to production.",
    size=11
)
add_bullet(doc, "STATIC_URL = 'static/'")
add_bullet(doc, "STATICFILES_STORAGE = 'whitenoise.storage.CompressedManifestStaticFilesStorage'")

doc.add_paragraph()
add_heading(doc, "9.2  Media Files", level=2, color_hex="2E6DA4")
add_paragraph(doc,
    "User-uploaded files (gallery images, event banners, teacher photos, etc.) "
    "are managed efficiently using cloud storage (e.g. Cloudinary) or served securely depending on infrastructure.",
    size=11
)
add_bullet(doc, "MEDIA_URL = 'media/'")
add_bullet(doc, "Sub-folders: gallery/, events/, activities/, students/, teachers/")
add_paragraph(doc,
    "For production environments with ephemeral storage (Render, Heroku, etc.), external object storage is strongly recommended to persist media files.",
    size=11, italic=True
)

add_page_break(doc)

# ══════════════════════════════════════════════
# 10. EMAIL NOTIFICATION SYSTEM
# ══════════════════════════════════════════════

add_heading(doc, "10. Email Notification System", level=1)
add_paragraph(doc,
    "The project includes an automated email notification system for key events:",
    size=11
)
doc.add_paragraph()

add_heading(doc, "10.1  Admission Application Email", level=2, color_hex="2E6DA4")
add_paragraph(doc,
    "When a new admission form is submitted, the admissions view uses Django's send_mail "
    "to dispatch an email to DEFAULT_FROM_EMAIL (noreply@littlesteps.com) containing "
    "the child name, parent details, phone, email, and chosen programme.",
    size=11
)

add_heading(doc, "10.2  Enquiry Email", level=2, color_hex="2E6DA4")
add_paragraph(doc,
    "When a contact form is submitted, the contact view sends a similar notification "
    "email with the enquirer's name, phone, email, and message.",
    size=11
)

add_heading(doc, "10.3  Bulk Parent Notification (Admin Action)", level=2, color_hex="2E6DA4")
add_paragraph(doc,
    "From the admin panel, staff can select one or more Notice/Event records and run "
    "the 'Send email notification to all parents' action. This collects all parent_email "
    "values from enrolled students and sends individual emails for each selected item.",
    size=11
)

doc.add_paragraph()
add_heading(doc, "10.4  Email Backend Configuration", level=2, color_hex="2E6DA4")
add_code_block(doc,
"""\
# settings.py (Production)
from decouple import config

EMAIL_BACKEND = 'django.core.mail.backends.smtp.EmailBackend'
EMAIL_HOST    = 'smtp.gmail.com'
EMAIL_PORT    = 587
EMAIL_USE_TLS = True
EMAIL_HOST_USER     = config('EMAIL_HOST_USER')
EMAIL_HOST_PASSWORD = config('EMAIL_HOST_PASSWORD')
DEFAULT_FROM_EMAIL  = 'noreply@littlestepsschool.com'\
""")

add_page_break(doc)

# ══════════════════════════════════════════════
# 11. INSTALLATION & SETUP GUIDE
# ══════════════════════════════════════════════

add_heading(doc, "11. Installation & Setup Guide", level=1)

steps = [
    ("Clone the repository",
     "git clone <repository-url>\ncd LittleSteps"),
    ("Create a virtual environment",
     "python -m venv venv\nvenv\\Scripts\\activate   # Windows\n# source venv/bin/activate  # macOS / Linux"),
    ("Setup Environment Variables",
     "Create a `.env` file in the root:\nDJANGO_SECRET_KEY=securekey\nDEBUG=False\nDB_NAME=littlesteps_db\nDB_USER=db_user\nDB_PASSWORD=secret"),
    ("Install dependencies",
     "pip install django pillow python-docx psycopg2-binary whitenoise python-decouple gunicorn"),
    ("Apply database migrations",
     "python manage.py makemigrations\npython manage.py migrate"),
    ("Create a superuser (admin account)",
     "python manage.py createsuperuser"),
    ("Collect static files (production only)",
     "python manage.py collectstatic"),
    ("Run the production server via Gunicorn",
     "gunicorn little_steps_web.wsgi:application --bind 0.0.0.0:8000"),
]

for i, (step_title, code) in enumerate(steps, 1):
    add_heading(doc, f"Step {i}: {step_title}", level=2, color_hex="2E6DA4")
    add_code_block(doc, code)
    doc.add_paragraph()

add_page_break(doc)

# ══════════════════════════════════════════════
# 12. RUNNING THE APPLICATION
# ══════════════════════════════════════════════

add_heading(doc, "12. Running the Application", level=1)

add_paragraph(doc,"After setup, access the portal at:", size=11)
add_bullet(doc, "Public site:    http://127.0.0.1:8000/")
add_bullet(doc, "Admin panel:    http://127.0.0.1:8000/admin/")
doc.add_paragraph()

add_heading(doc, "12.1  Admin Panel Usage", level=2, color_hex="2E6DA4")
admin_steps = [
    "Log in with the superuser credentials at /admin/.",
    "Use Notice to post school announcements.",
    "Use GalleryCategory + GalleryImage to manage the photo gallery.",
    "Use Event to add upcoming school events.",
    "Use ActivityUpdate to publish daily activity posts.",
    "Use VideoEmbed to add YouTube videos to the home page.",
    "Use AdmissionApplication to review online admission requests.",
    "Use Enquiry to read contact form submissions.",
    "Use Student to manage enrolled student records.",
    "Use Teacher to manage staff profiles displayed on the About page.",
]
for step in admin_steps:
    add_bullet(doc, step)

add_page_break(doc)

# ══════════════════════════════════════════════
# 13. FEATURES SUMMARY
# ══════════════════════════════════════════════

add_heading(doc, "13. Features Summary", level=1)

features = [
    ("Dynamic Home Page",          "Notices ticker, activity cards, upcoming events, embedded YouTube videos, testimonials"),
    ("About Page",                 "School history, mission & vision, programme offerings, teacher profiles"),
    ("Photo Gallery",              "Category-based image organisation with admin upload support"),
    ("Events Calendar",            "Automatic separation of upcoming vs past events"),
    ("Daily Updates / Blog",       "Activity update posts with images managed by admin"),
    ("Online Admission Form",      "Collects child & parent data, saves to DB, sends email notification via SMTP"),
    ("Contact / Enquiry Form",     "Saves enquiries to DB, sends email notification to school via SMTP"),
    ("Admin Back-end",             "Full CRUD for all models with filtering, search, and custom actions"),
    ("Bulk Email Notifications",   "Admin action to email all parents about notices or events"),
    ("Responsive Design",          "Modern CSS with mobile-friendly layouts"),
    ("Production-Ready",           "Utilizes PostgreSQL, Gunicorn, Whitenoise, and decouple for secrets"),
    ("Security Features",          "CSRF protection, X-Frame-Options, secure cookies, and disabled DEBUG"),
]

feat_tbl = doc.add_table(rows=1, cols=2)
feat_tbl.style = 'Table Grid'
hdr = feat_tbl.rows[0].cells
for i, h in enumerate(["Feature", "Description"]):
    hdr[i].text = h
    hdr[i].paragraphs[0].runs[0].bold = True
    set_cell_bg(hdr[i], "1F5C99")
    hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
for feature, desc in features:
    row = feat_tbl.add_row().cells
    row[0].text = feature
    row[1].text = desc

add_page_break(doc)

# ══════════════════════════════════════════════
# 14. FUTURE ENHANCEMENTS
# ══════════════════════════════════════════════

add_heading(doc, "14. Future Enhancements", level=1)

enhancements = [
    ("Parent Portal Login", "Allow parents to log in and track their child's admission status and activity feed."),
    ("Online Fee Payment",  "Integrate a payment gateway (Razorpay / Stripe) for school fee collection."),
    ("SMS Notifications",   "Extend notifications to SMS via Twilio or similar provider."),
    ("Progressive Web App", "Convert the portal to a PWA for offline access and push notifications."),
    ("Multi-language Support", "Add Hindi / regional language support for parents."),
    ("REST API",            "Expose a JSON API for a future React Native mobile app."),
    ("Cloud Media Storage", "Migrate media uploads to AWS S3 or Google Cloud Storage for scalability."),
    ("Analytics Dashboard", "Add visit and enquiry analytics for the school management."),
    ("Production Deployment","Deploy to PythonAnywhere / DigitalOcean with Gunicorn + Nginx + PostgreSQL."),
]

for title, desc in enhancements:
    add_bullet(doc, f"{desc}", bold_prefix=f"{title}: ")

doc.add_paragraph()
doc.add_paragraph()

# ══════════════════════════════════════════════
# FOOTER / Sign-off
# ══════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("─" * 60)
r.font.color.rgb = RGBColor(0xAA, 0xBB, 0xCC)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    f"Documentation generated on {datetime.date.today().strftime('%B %d, %Y')}  |  "
    "Little Steps International Pre School Web Portal v1.0"
)
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x99, 0xAA)
run.italic = True

# ──────────────────────────────────────────────
# Save the file
# ──────────────────────────────────────────────

output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Readme", "document.docx")
doc.save(output_path)
print(f"✅  Documentation saved → {output_path}")
